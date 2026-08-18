# -*- coding: utf-8 -*-
"""生成《赛菲莉娅背包整理插件 工作原理说明》Word 文档。"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, 'figures')
OUT = os.path.join(BASE, '赛菲莉娅背包整理插件-工作原理说明.docx')

DARK = RGBColor(0x1F, 0x3B, 0x57)
BLUE = RGBColor(0x2F, 0x6F, 0xAD)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_run(run, size=11, bold=False, color=None, name='微软雅黑'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)


def para(doc, text='', size=11, bold=False, color=None, align=None, space_after=6,
         space_before=0, indent=None, line=1.35):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if indent is not None:
        pf.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        set_run(r, size, bold, color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, 16, True, DARK)
    # 下边框线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '2F6FAD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, 13, True, BLUE)
    return p


def bullet(doc, text, size=11, space_after=3):
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run(r, size)
    return p


def num(doc, text, size=11, space_after=3):
    p = doc.add_paragraph(style='List Number')
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run(r, size)
    return p


def picture(doc, fname, width_in, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(os.path.join(FIG, fname), width=Inches(width_in))
    para(doc, caption, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths=None, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(htext)
        set_run(r, font, True, RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], '2F6FAD')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, font)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    para(doc, '', size=4, space_after=4)
    return t


doc = Document()

# 页面设置 A4
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.4)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(2.4)

# 默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ---------------- 封面 ----------------
for _ in range(4):
    para(doc, '', size=12)
para(doc, '赛菲莉娅 背包整理插件', size=30, bold=True, color=DARK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, '工作原理说明', size=26, bold=True, color=BLUE,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para(doc, '一键 F8 智能整理：识别、评分、搜索、应用', size=14, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
para(doc, 'BepInEx 插件 · v2.3.9 · Enhanced 增强整理模式', size=12, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
doc.add_page_break()

# ---------------- 1 ----------------
h1(doc, '1. 这个插件是干什么的')
para(doc, '《赛菲莉娅》的背包不是普通格子。石板会改变周围格子的等级；护符大多带位置条件'
          '（必须在顶行、必须在两侧、旁边必须空着……）；还有行星望远镜、和谐之晶、奉献徽章、'
          '指北针这类靠"摆在一起"才生效的组合机制。要同时满足所有这些条件，手动摆基本不可能，'
          '游戏自带的自动排列又只跑很少的迭代次数，经常停在明显不是最优的结果上。')
para(doc, '这个插件做的事很简单：按一下 F8，把背包重排一遍，让各种加成尽量同时吃到，排完提示'
          '"整理完毕"。它是挂进游戏进程的 BepInEx 插件，不修改任何游戏文件。整理分三种情况：')
bullet(doc, 'Vanilla 模式：直接调用游戏自带的自动排列（迭代次数可调）；')
bullet(doc, 'Enhanced 模式（默认）：插件自己计算最优布局再应用，是本文的重点；')
bullet(doc, '联机客户端：服务器不让客户端直接改背包，插件改用网络命令逐步调整，同样能吃到智能整理。')

# ---------------- 2 ----------------
h1(doc, '2. 总体思路：先离线算，再一次性应用')
para(doc, 'Enhanced 的核心做法是"离线评分模型"。按 F8 之后，插件先把背包当前状态完整复制一份，'
          '然后在副本上做三件事：')
num(doc, '识别：把每件物品归类，搞清楚哪些是石板、哪些是护符、各自带什么特殊机制；')
num(doc, '评分：给"任意一种摆法"打一个分，分数越高，游戏里的实际加成越高；')
num(doc, '搜索：从智能初始布局出发，跑多轮模拟退火，找出分数最高的布局。')
para(doc, '整个过程不碰游戏状态，最后把最优布局一次性写回游戏。这样设计有三个好处：')
bullet(doc, '可以在副本上跑几千次迭代。游戏内置算法迭代次数很少（原版默认仅 4 次），离线副本没有这个顾虑，搜得更彻底；')
bullet(doc, '游戏里只改一次。计算在后台瞬间完成，写回也是一下子的事，不会出现"看着物品自己挪来挪去"的等待；')
bullet(doc, '联机客户端能用。没有服务器权限也能算，算完用网络命令执行（见第 6 节）。')
picture(doc, 'fig_overview.png', 6.3, '图 1  单次按 F8 的完整处理流程（总流程）')

# ---------------- 3 ----------------
h1(doc, '3. 第一步：认识背包里的每一件东西')
para(doc, '背包物品先分成四类：')
bullet(doc, '石板：占多格，覆盖一个效果网格（给范围内的格子加等级、乘倍率，或解锁"豁免格"），部分石板可旋转；')
bullet(doc, '护符：占一格，有等级，大多数带位置条件或联动机制；')
bullet(doc, '负面藏品（负担）：像心之重担，放哪里都扣分，只能尽量塞进最差的格子减少损失；')
bullet(doc, '杂物：凑数物品，填空用。')
para(doc, '插件的识别大部分靠类型或接口判断（不依赖物品显示名），少数机制物品用配置里的 key 兜底。'
          '全部特殊机制见下表：')
table(doc,
      ['机制', '识别方式', '对整理的意义'],
      [
          ['位置条件护符', 'CharmPositionKind 枚举（顶行/底行/两侧/内侧/外侧/两侧空/两侧有护符/八邻域满/靠近魔法书）',
           '不满足条件的护符会被游戏禁用，优先放满足条件的格'],
          ['豁免格', 'IgnoreCriteria 解锁石板产生的格子', '放上去无视位置条件，冰锁优先站豁免格'],
          ['行星望远镜', 'Charm_PlanetModule 类型', '周围 8 格每颗启用行星都有加成'],
          ['行星藏品', '分类标签 PLANET', '聚到望远镜旁边；可配置排除乐谱银河、红色行星观察日志'],
          ['和谐之晶', 'Charm_NearLevelDamage 类型', '周围 8 格护符等级和越高，伤害放大越多'],
          ['奉献徽章', 'Charm_CompanionChaos 类型 + 配置 key', '同一横排的同伴藏品全部进入强化态'],
          ['同伴藏品', 'ICompanionCharm 接口', '金色手铃/迷你弩炮/灵魂粉末/采矿臂章等自动识别'],
          ['指北针', 'Charm_UpCharmDamage 类型', '上方是伤害类藏品或另一块指北针才生效'],
          ['行锁定物品', '配置 key（凯尔萨德尼钥匙）', '藏品类型随所在行变化，整理绝不跨行'],
          ['神秘地块', '分类标签 Mystic', '凑齐 2 个→1 格等级×2，5 个→4 格×2'],
          ['武器相关护符', 'isWeaponRelatedCharm + 当前武器', '武器类型匹配才启用（如丢弃的金戒指）'],
      ],
      widths=[3.0, 6.2, 6.6])
para(doc, '值得说明的两点：豁免格是某些石板的效果，冰锁这类护符（默认配置里指定）会优先站豁免格，'
          '因为豁免格往往等级更高；行锁定物品（凯尔萨德尼钥匙）的类型随所在行变化，插件只在它所在行内'
          '调整位置，绝不变行。下面这张图是"一张布局能同时吃到的机制"示意——每种颜色代表一类机制：', space_after=4)
picture(doc, 'fig_backpack.png', 5.9, '图 2  一张布局能同时吃到的机制（示意）')

# ---------------- 4 ----------------
h1(doc, '4. 第二步：给一种摆法打分')
para(doc, '评分函数是插件的"世界观"：它离线模拟游戏里所有的加成公式，任何摆法都算出一个分数，'
          '后面的搜索就以这个分数为目标。')
h2(doc, '4.1 先算每个格子的有效等级')
para(doc, '格子等级由三部分叠加：石板效果网格的加成（+N 或 ×M）、护符自带的附魔（+N）、神秘地块的'
          '倍率（×2）。插件把"物品拿走后格子本身的裸等级"作为底数，再把石板贡献、附魔、神秘倍率'
          '依次叠上去，得到当前摆法下每个格子的有效等级。')
h2(doc, '4.2 逐格判断护符是否启用')
para(doc, '按游戏规则检查每件护符：有效等级 ≥ 0、位置条件满足（或站在豁免格）、武器匹配（武器相关'
          '护符）。启用与禁用的分差是 1750（+1000 对 -750），这一项直接把"摆对了"和"摆错了"拉开差距；'
          '站上负等级格子还要额外扣分（每级 -250）。')
h2(doc, '4.3 再叠加各机制的分')
para(doc, '评分里的每一项都对应游戏里真实存在的机制，数值是权重，代表插件认为这个机制值多少分。'
          '各项默认值见图 3 与表 2。')
table(doc,
      ['评分项', '默认权重', '说明'],
      [
          ['护符等级分', '有效等级 × 10000 × 优先级权重', '启用的护符按有效等级计分，优先级越高权重越大'],
          ['启用/禁用', '+1000 / -750', '摆对位置与摆错位置的基础分差'],
          ['负等级格子', '-250 / 级', '护符站上负等级格子的惩罚'],
          ['优先级', 'P1=1.5 / P2=1.25 / P3=1.1 / P4=1.0', '传说/羁绊=1，稀有=2，高级=3，普通=4，可配置'],
          ['位置条件', '+500（冰锁豁免格 +5000）', '引导受限护符站到满足条件的格子'],
          ['行星聚簇', '+40000 / 颗', '望远镜周围每颗启用行星，远高于行星自身等级分，搜索会优先聚拢'],
          ['和谐之晶', '+2000 × 周围8格等级和', '等级和越高伤害放大越多，高等级护符被吸引到晶周围'],
          ['奉献徽章', '+3000 / 个同行同伴', '徽章须启用；同伴自动按 ICompanionCharm 识别'],
          ['指北针配对', '+12000 × 上方优先级权重', '上方是伤害类/指北针才给分；未配对等级分只算一成'],
          ['负担惩罚', '-20000 / 级高出', '负担必须待在最低等级格，否则按高出级数扣分'],
          ['行锁定', '-100000 / 件', '跨行强约束，退火几乎不可能接受跨行布局'],
          ['神秘地块', '等级直接 ×2', '高价值护符优先放到 ×2 地块上'],
      ],
      widths=[3.4, 5.6, 6.8])
picture(doc, 'fig_weights.png', 6.3, '图 3  评分模型各项权重（默认配置）')

h2(doc, '4.4 等级是谁"分配"的')
para(doc, '前面算分用到"格子等级"，这里把它的来源讲清楚：等级不是物品自带的属性，而是格子的属性。'
          '格子的有效等级由三部分组成：')
bullet(doc, '裸等级（baseLevel）：把当前所有石板效果、物品附魔、神秘倍率剥掉之后，格子固有的等级，来自游戏的 levelMatrix；')
bullet(doc, '石板贡献：石板效果网格（+N 等级 / ×M 倍率 / 禁用 / 解锁豁免格）盖在哪些格子上；')
bullet(doc, '附魔：物品自带的等级，直接加在所在格上。')
para(doc, '所以真正在"分配等级"的是石板的位置和旋转。打个比方：石板是聚光灯，决定哪几格亮、亮多少、'
          '哪几格禁用或豁免；护符是演员，带位置条件的只能站特定区域，其余的按优先级挑亮的格子站。')
para(doc, '"先摆石板后放物品"不是死板的先后，而是交替推进：每摆一块石板（或每放一件物品）后，插件都会'
          '按当前布局重新计算一遍所有格子的等级，下一步的决策基于最新的等级。石板摆位时能看到护符将来'
          '站哪（打分函数会检查已占用的格子），护符选格时能看到石板把哪里加亮了。')
para(doc, '石板摆位本身也在为护符铺路：负效果石板最先摆，评分会额外惩罚"负等级效果落在空格上"，'
          '避免负等级漏到将来放护符的格子上；豁免格、禁用格也按各自权重计分。')
para(doc, '最终保证"恰到好处"的是搜索阶段：退火里约 18% 的变异是旋转石板（改灯光方向），其余是移动、'
          '交换物品（改演员站位），每次变异都重算等级再打分，直到评分收敛。所以最终布局是石板与物品'
          '联合优化出来的，而不是"先定石板、物品听天由命"。')
picture(doc, 'fig_levels.png', 6.4, '图 4  等级分配：裸等级地形 → 石板打光 → 护符就位')

# ---------------- 5 ----------------
h1(doc, '5. 第三步：怎么找到最高分的布局')
para(doc, '搜索分两个阶段：先搭一个不错的起点（智能初始布局），再用模拟退火在这个起点附近深挖。')
h2(doc, '5.1 智能初始布局')
para(doc, '纯随机起点收敛慢，所以先按规则搭一个起点，顺序见图 4：')
num(doc, '石板先摆：负效果的石板优先，逐格逐旋转打分，选覆盖最好的位置；')
num(doc, '受限护符先放：顺序是"满足条件的格 → 豁免格 → 任意格"；')
num(doc, '行锁定物品放回用户所在行，行内选最佳列；')
num(doc, '行星望远镜落位，作为行星聚簇的锚点；')
num(doc, '和谐之晶落位，记录它周围 8 格（这些格加权，吸引高等级护符）；')
num(doc, '奉献徽章落位，记录它所在的行（同行格子加权，引导同伴同排）；')
num(doc, '行星聚到望远镜相邻格（配置排除的行星类藏品除外）；')
num(doc, '指北针优先找"上方是伤害类藏品"的格子；')
num(doc, '其余护符按用户优先级 P1→P4，同优先级内按稀有度；')
num(doc, '杂物填空，负担塞最差（负等级最高）的格子。')
para(doc, '给单个护符选格时，选格函数会把等级、位置条件、豁免格、和谐之晶邻域（+8000）、奉献徽章同行'
          '（+6000）、神秘×2 地块（+20000）这些加权全部算进去，一次选到当前最合适的格。')
picture(doc, 'fig_smartstart.png', 5.6, '图 5  智能初始布局：按机制优先级逐个落位')
h2(doc, '5.2 模拟退火')
para(doc, '起点只是一个不错的解，不一定最优。退火阶段反复做微小改动，每改一次重新打分：')
bullet(doc, '约 38% 的改动是"定向移动"：把受限护符直接跳到满足条件的格子、把行星挪到望远镜旁、'
            '把指北针挪到伤害类下方、把负担丢进负格（图 5）；')
bullet(doc, '其余是随机探索：随机移动、随机交换两个格子、随机旋转石板。')
para(doc, '关键在温度：一开始温度高，分数变差也大概率接受，这样能翻过"小山坡"，不会困在局部最优里；'
          '温度随时间线性降到 1，后期只接受变好的结果，保证收敛（图 6）。每轮退火还会随机重启 3 次，'
          '每次从当前最优继续。')
picture(doc, 'fig_mutation.png', 4.6, '图 6  每次迭代的变异操作构成：定向移动 + 随机探索')
picture(doc, 'fig_anneal.png', 6.3, '图 7  模拟退火的温度曲线与接受概率')
h2(doc, '5.3 多轮独立搜索')
para(doc, '每次按 F8，插件内部用 4 个不同随机种子各跑一遍完整搜索，取全局最高分。等效于自动把 F8 '
          '重复按了 4 次，一次到位。因为全程离线评估只有毫秒级，34 格满包一次整理约 200~300ms，'
          '游戏里感觉不到卡顿。')

# ---------------- 6 ----------------
h1(doc, '6. 第四步：把结果写回游戏')
para(doc, '单机或主机：直接把算好的整包布局写回背包，再用游戏自身接口读一次真实评分核对，日志里对比'
          '"离线评分 / 游戏评分"。')
para(doc, '联机客户端（图 7）：客户端没有服务器权限，插件把"当前布局 → 目标布局"翻译成两类操作序列：')
bullet(doc, '交换：任意两格物品互换（Mirror 的 Swap 命令，需要服务器授权）；')
bullet(doc, '旋转：把某块石板点击旋转若干次（DoClickAction）。')
para(doc, '转换时在内存里推演每件物品的位置与旋转，生成最短的交换/旋转序列，再逐条执行。整个过程'
          '不清空背包，比主机版更保守——这也是早期版本踩过丢物品的坑之后换来的经验（见第 7 节）。')
picture(doc, 'fig_client.png', 6.3, '图 8  联机客户端整理：离线算最优 → 生成操作序列 → 网络命令执行')

# ---------------- 7 ----------------
h1(doc, '7. 安全设计：宁可不动，不可丢物品')
para(doc, '整理是重排整个背包，一旦在错误时机动手就可能丢东西。插件有几层保护：')
bullet(doc, '会话稳定延迟：进图后等 3 秒（SessionStableDelay）。联机房间刚进时背包初始化没完成，'
            '此刻按 F8 会提示"请稍候"而不是动手；')
bullet(doc, '快照校验：整理前先抓一份背包快照，只统计正常背包格（排除药水带这类特殊区域），'
            '与游戏当前状态对不上就取消本次整理；')
bullet(doc, '一次应用：算完只应用一次，失败不重试，宁可留着不整理也不冒险。')
para(doc, '这层防御是有代价换来的：早期版本用"清空再重写"的方式应用布局，触发了游戏的起始物品补货'
          '机制，出现过物品丢失/复制的问题。后来改成现在的防御式设计。')

# ---------------- 8 ----------------
h1(doc, '8. 实测效果')
bullet(doc, '一次 F8 直达最佳：多轮搜索合并后，单次按键即达到此前反复按多次的效果；')
bullet(doc, '和谐之晶实测：3 块晶互聚成簇，14 级护符放在三者公共邻域格，每个晶周围 8 格等级和 = 16，'
            '离线评分 4000 → 355004（+351004），耗时 141ms；')
bullet(doc, '每次整理日志会输出：识别统计（护符/望远镜/罗盘/附魔/奉献徽章/同伴数量）、布局网格图、'
            '各机制落地位置（"望远镜@(x,y)：相邻行星 N 颗""奉献徽章@(x,y)：同行同伴 N 个"等），'
            '方便核对插件把东西摆到了哪里。')

# ---------------- 9 ----------------
h1(doc, '9. 配置速查')
para(doc, '所有参数都在 BepInEx 配置文件里（com.sephiria.backpack-organizer.cfg），改完重启游戏生效。'
          '常用项见下表：')
table(doc,
      ['分区 / 配置项', '默认值', '作用'],
      [
          ['General / Hotkey', 'F8', '触发整理的快捷键'],
          ['General / SortMode', 'Enhanced', 'Vanilla（游戏内置）/ Enhanced（增强）'],
          ['General / SessionStableDelay', '3 秒', '进图后等待背包初始化完成的秒数（防丢物品）'],
          ['General / RowLockedItems', '凯尔萨德尼钥匙 key', '行锁定物品：保持所在行不变'],
          ['Enhanced / Iterations', '3000', '模拟退火迭代次数（离线评估，可放心调大）'],
          ['Enhanced / Restarts', '3', '每轮退火的随机重启次数'],
          ['Enhanced / Temperature', '800', '初始温度：越大越容易跳出局部最优'],
          ['Enhanced / SearchRounds', '4', '每按一次 F8 跑的独立搜索轮数'],
          ['Smart / EnableSmartStart', 'true', '智能初始布局开关'],
          ['Smart / EnableRandomStarts', 'true', '多起点搜索（智能初始/原始/随机）'],
          ['Priority / Enable', 'true', '优先级系统：传说/羁绊=1，稀有=2，高级=3，普通=4'],
          ['Priority / FixedHighPriorityItems', '冰锁/金戒指/绝对戒指/红茶叶袋', '强制最高优先级的特定物品'],
          ['Priority / Weight1~4', '1.5/1.25/1.1/1.0', '各优先级护符的等级分权重'],
          ['Synergy / PlanetBonus', '40000', '行星聚簇：望远镜周围每颗行星'],
          ['Synergy / HarmonyLevelBonus', '2000', '和谐之晶：周围 8 格每级护符等级'],
          ['Synergy / DedicationCompanionBonus', '3000', '奉献徽章：每个同行同伴'],
          ['Synergy / CompassBonus', '12000', '指北针配对奖励'],
          ['Burden / NegativeCellPenalty', '20000', '负面藏品未待负格时的扣分'],
          ['Mystic / Enable', 'true', '神秘 ×2 地块联动'],
          ['Debug / SelfTest', 'false', '自检模式：打乱后整理，对比前后评分'],
      ],
      widths=[5.4, 4.6, 5.8])

# ---------------- 10 ----------------
h1(doc, '10. 附：版本演进')
table(doc,
      ['版本', '关键变化'],
      [
          ['v1.x', '调用游戏内置自动排列，仅主机可用'],
          ['v2.0', '全离线评分模型 + 模拟退火，迭代次数不受帧循环限制'],
          ['v2.1', '附魔计入、神秘 ×2 地块（宝珠功能因复制 bug 移除）'],
          ['v2.2', '指北针配对、负担塞负格、豁免格优先级'],
          ['v2.3', '优先级系统、行星聚簇（含排除项）、行锁定物品、多轮搜索一次到位'],
          ['v2.3.6', '武器相关护符完整武器匹配（修复丢弃的金戒指拿不到加成）'],
          ['v2.3.8', '和谐之晶：周围 8 格等级和放大，聚簇引导'],
          ['v2.3.9', '奉献徽章：同一横排同伴聚拢（本次新增）'],
      ],
      widths=[3.0, 12.8])

para(doc, '', size=8)
para(doc, '本文档随插件 v2.3.9 发布，机制说明以当前版本代码为准。', size=9, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print('docx saved ->', OUT)
